"""
Document processing utilities for medical data ingestion
"""
import os
import json
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import PyPDF2
from docx import Document
import requests
from bs4 import BeautifulSoup


class MedicalDocument:
    """Represents a medical document with metadata"""
    
    def __init__(
        self,
        content: str,
        title: str,
        source: str,
        doc_id: Optional[str] = None,
        year: Optional[str] = None,
        url: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ):
        self.content = content
        self.title = title
        self.source = source
        self.doc_id = doc_id or self._generate_doc_id(content, title)
        self.year = year
        self.url = url
        self.metadata = metadata or {}
        self.processed_at = datetime.now().isoformat()
    
    def _generate_doc_id(self, content: str, title: str) -> str:
        """Generate unique document ID from content and title"""
        hash_input = f"{title}:{content[:500]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary"""
        return {
            "doc_id": self.doc_id,
            "content": self.content,
            "title": self.title,
            "source": self.source,
            "year": self.year,
            "url": self.url,
            "metadata": self.metadata,
            "processed_at": self.processed_at
        }


class DocumentLoader:
    """Handles loading documents from various sources"""
    
    @staticmethod
    def load_pdf(file_path: str) -> List[str]:
        """Extract text from PDF file"""
        text_chunks = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_chunks.append(text)
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
        return text_chunks
    
    @staticmethod
    def load_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def load_txt(file_path: str) -> str:
        """Load text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
            return ""
    
    @staticmethod
    def load_json(file_path: str) -> List[Dict[str, Any]]:
        """Load structured data from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except Exception as e:
            print(f"Error loading JSON {file_path}: {e}")
            return []


class PubMedLoader:
    """Fetch medical abstracts from PubMed"""
    
    BASE_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    
    def __init__(self, email: str = "your_email@example.com"):
        self.email = email
    
    def search_pubmed(self, query: str, max_results: int = 100) -> List[str]:
        """Search PubMed and return list of PMIDs"""
        search_url = f"{self.BASE_URL}esearch.fcgi"
        params = {
            "db": "pubmed",
            "term": query,
            "retmax": max_results,
            "retmode": "json",
            "email": self.email
        }
        
        try:
            response = requests.get(search_url, params=params)
            response.raise_for_status()
            data = response.json()
            return data.get("esearchresult", {}).get("idlist", [])
        except Exception as e:
            print(f"Error searching PubMed: {e}")
            return []
    
    def fetch_abstracts(self, pmids: List[str]) -> List[MedicalDocument]:
        """Fetch abstracts for given PMIDs"""
        if not pmids:
            return []
        
        fetch_url = f"{self.BASE_URL}efetch.fcgi"
        params = {
            "db": "pubmed",
            "id": ",".join(pmids),
            "retmode": "xml",
            "email": self.email
        }
        
        documents = []
        try:
            response = requests.get(fetch_url, params=params)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'xml')
            
            articles = soup.find_all('PubmedArticle')
            for article in articles:
                try:
                    pmid = article.find('PMID').text
                    title_elem = article.find('ArticleTitle')
                    title = title_elem.text if title_elem else "No title"
                    
                    abstract_elem = article.find('AbstractText')
                    abstract = abstract_elem.text if abstract_elem else ""
                    
                    year_elem = article.find('PubDate')
                    year = None
                    if year_elem and year_elem.find('Year'):
                        year = year_elem.find('Year').text
                    
                    if abstract:
                        doc = MedicalDocument(
                            content=abstract,
                            title=title,
                            source="PubMed",
                            doc_id=f"PMID_{pmid}",
                            year=year,
                            url=f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                            metadata={"pmid": pmid}
                        )
                        documents.append(doc)
                except Exception as e:
                    print(f"Error parsing article: {e}")
                    continue
        
        except Exception as e:
            print(f"Error fetching abstracts: {e}")
        
        return documents


class MedicalCorpusBuilder:
    """Build and manage medical document corpus"""
    
    def __init__(self, data_dir: str = "./data"):
        self.data_dir = Path(data_dir)
        self.raw_dir = self.data_dir / "raw_documents"
        self.processed_dir = self.data_dir / "processed"
        
        # Create directories
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
    
    def load_local_documents(self) -> List[MedicalDocument]:
        """Load documents from local directory"""
        documents = []
        loader = DocumentLoader()
        
        for file_path in self.raw_dir.rglob("*"):
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                
                try:
                    if suffix == ".pdf":
                        pages = loader.load_pdf(str(file_path))
                        for i, page_text in enumerate(pages):
                            doc = MedicalDocument(
                                content=page_text,
                                title=f"{file_path.stem} - Page {i+1}",
                                source=file_path.name,
                                metadata={"file_path": str(file_path), "page": i+1}
                            )
                            documents.append(doc)
                    
                    elif suffix == ".docx":
                        content = loader.load_docx(str(file_path))
                        if content:
                            doc = MedicalDocument(
                                content=content,
                                title=file_path.stem,
                                source=file_path.name,
                                metadata={"file_path": str(file_path)}
                            )
                            documents.append(doc)
                    
                    elif suffix == ".txt":
                        content = loader.load_txt(str(file_path))
                        if content:
                            doc = MedicalDocument(
                                content=content,
                                title=file_path.stem,
                                source=file_path.name,
                                metadata={"file_path": str(file_path)}
                            )
                            documents.append(doc)
                    
                    elif suffix == ".json":
                        json_data = loader.load_json(str(file_path))
                        for item in json_data:
                            if isinstance(item, dict) and "content" in item:
                                doc = MedicalDocument(
                                    content=item.get("content", ""),
                                    title=item.get("title", file_path.stem),
                                    source=item.get("source", file_path.name),
                                    year=item.get("year"),
                                    url=item.get("url"),
                                    metadata=item.get("metadata", {})
                                )
                                documents.append(doc)
                
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
        
        return documents
    
    def save_corpus(self, documents: List[MedicalDocument], filename: str = "corpus.json"):
        """Save processed corpus to JSON file"""
        corpus_path = self.processed_dir / filename
        corpus_data = [doc.to_dict() for doc in documents]
        
        with open(corpus_path, 'w', encoding='utf-8') as f:
            json.dump(corpus_data, f, indent=2, ensure_ascii=False)
        
        print(f"Saved {len(documents)} documents to {corpus_path}")
        return corpus_path
    
    def load_corpus(self, filename: str = "corpus.json") -> List[MedicalDocument]:
        """Load processed corpus from JSON file"""
        corpus_path = self.processed_dir / filename
        
        if not corpus_path.exists():
            return []
        
        with open(corpus_path, 'r', encoding='utf-8') as f:
            corpus_data = json.load(f)
        
        documents = []
        for data in corpus_data:
            doc = MedicalDocument(
                content=data["content"],
                title=data["title"],
                source=data["source"],
                doc_id=data["doc_id"],
                year=data.get("year"),
                url=data.get("url"),
                metadata=data.get("metadata", {})
            )
            documents.append(doc)
        
        return documents
